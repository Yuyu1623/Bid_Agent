# 这个脚本负责把 PDF、Word 等招标文件解析成分章节内容。
# -*- coding: utf-8 -*-
import os
import re
import zipfile
import importlib.util
from contextlib import contextmanager
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any, List, Optional

from pydantic import BaseModel, Field

from MinerU_pdf_parse_tool import mineru_parse_document


class BidDocumentSection(BaseModel):
    index: int = Field(description="Section order in the parsed bid document.")
    title: str = Field(description="Detected section title.")
    level: int = Field(description="Heading level. Smaller numbers are higher level.")
    content: str = Field(description="Section body without the heading line.")
    markdown: str = Field(description="Section heading and body in Markdown.")
    metadata: dict[str, Any] = Field(
        default_factory=dict,
        description="Structured metadata such as title path, tables, images and page hints.",
    )


MINERU_PARSE_METHODS = {"mineru_vlm", "mineru_pipeline", "mineru_html"}
PARALLEL_MINERU_PARSE_METHODS = {"mineru_parallel_pages"}
LOCAL_PARSE_METHODS = {
    "pymupdf4llm",
    "docling",
    "pdfplumber",
    "docx2python",
    "docx2python_image_ocr",
}
SUPPORTED_PARSE_METHODS = (
    MINERU_PARSE_METHODS
    | PARALLEL_MINERU_PARSE_METHODS
    | LOCAL_PARSE_METHODS
)

MARKDOWN_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
CHINESE_CHAPTER_RE = re.compile(
    r"^\s*(第[一二三四五六七八九十百千万零〇\d]+[章节篇部分][\s、：:.-]*.+?)\s*$"
)
CHINESE_ORDER_RE = re.compile(r"^\s*([一二三四五六七八九十百千万零〇]+[、.．]\s*.+?)\s*$")
NUMBERED_HEADING_RE = re.compile(r"^\s*((?:\d+[.．、]){1,5}\s*.+?)\s*$")
PAREN_HEADING_RE = re.compile(r"^\s*([（(][一二三四五六七八九十百千万零〇\d]+[）)]\s*.+?)\s*$")


def parse_bid_document(
    document: str,
    output_dir: Optional[str] = None,
    parse_method: str = "mineru_vlm",
    model_version: Optional[str] = None,
    language: str = "ch",
    is_ocr: Optional[bool] = None,
    enable_formula: bool = True,
    enable_table: bool = True,
    page_ranges: Optional[str] = None,
    poll_interval: int = 5,
    timeout: int = 600,
) -> List[BidDocumentSection]:
    """Parse a bid document and split it into section objects."""
    markdown = parse_document_to_markdown(
        document=document,
        output_dir=output_dir,
        parse_method=parse_method,
        model_version=model_version,
        language=language,
        is_ocr=is_ocr,
        enable_formula=enable_formula,
        enable_table=enable_table,
        page_ranges=page_ranges,
        poll_interval=poll_interval,
        timeout=timeout,
    )
    sections = split_bid_markdown_sections(markdown)
    return enrich_sections_with_metadata(sections, document=document)


def parse_document_to_markdown(
    document: str,
    output_dir: Optional[str] = None,
    parse_method: str = "mineru_vlm",
    model_version: Optional[str] = None,
    language: str = "ch",
    is_ocr: Optional[bool] = None,
    enable_formula: bool = True,
    enable_table: bool = True,
    page_ranges: Optional[str] = None,
    poll_interval: int = 5,
    timeout: int = 600,
) -> str:
    """Parse a local/remote document to Markdown with MinerU or a local parser."""
    parse_method = _normalize_parse_method(document, parse_method)
    if parse_method not in SUPPORTED_PARSE_METHODS:
        supported = ", ".join(sorted(SUPPORTED_PARSE_METHODS))
        raise ValueError(f"Unsupported parse_method: {parse_method}. Supported: {supported}")

    if parse_method in MINERU_PARSE_METHODS:
        mineru_model_version, mineru_is_ocr = _resolve_mineru_method(
            parse_method=parse_method,
            model_version=model_version,
            is_ocr=is_ocr,
        )
        mineru_enable_formula = enable_formula if mineru_model_version != "MinerU-HTML" else False
        mineru_enable_table = enable_table if mineru_model_version != "MinerU-HTML" else False
        return mineru_parse_document(
            document=document,
            output_dir=output_dir,
            model_version=mineru_model_version,
            language=language,
            is_ocr=mineru_is_ocr,
            enable_formula=mineru_enable_formula,
            enable_table=mineru_enable_table,
            page_ranges=page_ranges,
            poll_interval=poll_interval,
            timeout=timeout,
        )

    if parse_method == "mineru_parallel_pages":
        return _parse_with_parallel_mineru_pages(
            document=document,
            output_dir=output_dir,
            model_version=model_version or "vlm",
            language=language,
            is_ocr=bool(is_ocr) if is_ocr is not None else False,
            enable_formula=enable_formula,
            enable_table=enable_table,
            page_ranges=page_ranges,
            poll_interval=poll_interval,
            timeout=timeout,
        )

    if parse_method == "pymupdf4llm":
        return _parse_pdf_with_pymupdf4llm(document=document, page_ranges=page_ranges)
    if parse_method == "docling":
        return _parse_pdf_with_docling(document=document, page_ranges=page_ranges)
    if parse_method == "pdfplumber":
        return _parse_pdf_with_pdfplumber(document=document, page_ranges=page_ranges)
    if parse_method == "docx2python":
        return _parse_docx_with_docx2python(document=document, include_image_ocr=False)
    if parse_method == "docx2python_image_ocr":
        return _parse_docx_with_docx2python(document=document, include_image_ocr=True)

    raise ValueError(f"Unsupported parse_method: {parse_method}")


def _normalize_parse_method(document: str, parse_method: Optional[str]) -> str:
    """Convert UI/API aliases such as auto into a concrete parser name."""
    method = (parse_method or "mineru_vlm").strip()
    if method != "auto":
        return method

    path = Path(document)
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            with zipfile.ZipFile(path) as archive:
                image_count = sum(
                    1
                    for name in archive.namelist()
                    if name.startswith("word/media/") and not name.endswith("/")
                )
        except Exception:
            image_count = 0
        return "docx2python_image_ocr" if image_count >= 8 else "docx2python"
    if suffix == ".doc":
        return "docx2python_image_ocr"
    if suffix == ".pdf":
        if not _pdf_has_text_layer(path):
            page_count = _get_pdf_page_count(path)
            return "mineru_parallel_pages" if page_count >= 30 else "mineru_vlm"
        if _pdf_has_images(path):
            return "mineru_parallel_pages"
        if _pdf_is_simple_text_file(path):
            return "pdfplumber"
        page_count = _get_pdf_page_count(path)
        return "mineru_parallel_pages" if page_count >= 30 else "mineru_pipeline"
    if suffix in {".png", ".jpg", ".jpeg", ".webp"}:
        return "mineru_vlm"
    return "mineru_vlm"


def _pdf_is_simple_text_file(path: Path) -> bool:
    """Return True only for lightweight native PDFs without obvious layout risk."""
    if path.suffix.lower() != ".pdf" or not path.exists():
        return False
    try:
        import pdfplumber
    except ImportError:
        return False

    try:
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            if page_count > int(os.getenv("PDFPLUMBER_SIMPLE_MAX_PAGES", "20")):
                return False
            selected_pages = pdf.pages[: min(page_count, 5)]
            if not selected_pages:
                return False
            total_chars = 0
            table_pages = 0
            image_count = 0
            for page in selected_pages:
                text = page.extract_text() or ""
                total_chars += len(text.strip())
                image_count += len(getattr(page, "images", []) or [])
                try:
                    if page.extract_tables():
                        table_pages += 1
                except Exception:
                    pass
            chars_per_page = total_chars / max(1, len(selected_pages))
            return image_count == 0 and table_pages == 0 and chars_per_page >= 300
    except Exception:
        return False


def _parse_with_parallel_mineru_pages(
    document: str,
    output_dir: Optional[str],
    model_version: str,
    language: str,
    is_ocr: bool,
    enable_formula: bool,
    enable_table: bool,
    page_ranges: Optional[str],
    poll_interval: int,
    timeout: int,
) -> str:
    path = Path(document).expanduser().resolve()
    if path.suffix.lower() != ".pdf":
        return mineru_parse_document(
            document=document,
            output_dir=output_dir,
            model_version=model_version,
            language=language,
            is_ocr=is_ocr,
            enable_formula=enable_formula,
            enable_table=enable_table,
            page_ranges=page_ranges,
            poll_interval=poll_interval,
            timeout=timeout,
        )

    page_count = _get_pdf_page_count(path)
    if page_count <= 0:
        raise RuntimeError("Cannot determine PDF page count for parallel MinerU parsing.")

    selected_pages = sorted(_parse_page_ranges(page_ranges) or set(range(1, page_count + 1)))
    out_of_range_pages = [page for page in selected_pages if page < 1 or page > page_count]
    if out_of_range_pages:
        raise RuntimeError(
            f"Page range contains pages outside the PDF page count ({page_count}): "
            f"{_format_page_range(out_of_range_pages)}"
        )
    if not selected_pages:
        raise RuntimeError("No pages selected for parallel MinerU parsing.")

    chunk_size = max(1, int(os.getenv("MINERU_PARALLEL_PAGE_CHUNK_SIZE", "30")))
    max_workers = max(1, int(os.getenv("MINERU_PARALLEL_MAX_WORKERS", "3")))
    page_chunks = _chunk_contiguous_pages(selected_pages, chunk_size=chunk_size)

    with TemporaryDirectory() as chunk_dir:
        chunk_jobs = []
        for index, pages in enumerate(page_chunks, start=1):
            page_range = _format_page_range(pages)
            chunk_path = Path(chunk_dir) / f"pages_{page_range.replace('-', '_')}.pdf"
            _write_pdf_page_chunk(source_pdf=path, pages=pages, output_pdf=chunk_path)
            chunk_jobs.append((index, page_range, chunk_path))

        if len(chunk_jobs) == 1:
            _, page_range, chunk_path = chunk_jobs[0]
            try:
                return mineru_parse_document(
                    document=str(chunk_path),
                    output_dir=output_dir,
                    model_version=model_version,
                    language=language,
                    is_ocr=is_ocr,
                    enable_formula=enable_formula,
                    enable_table=enable_table,
                    page_ranges=None,
                    poll_interval=poll_interval,
                    timeout=timeout,
                )
            except Exception as exc:
                raise RuntimeError(f"MinerU parallel page chunk {page_range} failed: {exc}") from exc

        def run_chunk(index: int, page_range: str, chunk_path: Path) -> tuple[int, str, str]:
            try:
                markdown = mineru_parse_document(
                    document=str(chunk_path),
                    output_dir=output_dir,
                    model_version=model_version,
                    language=language,
                    is_ocr=is_ocr,
                    enable_formula=enable_formula,
                    enable_table=enable_table,
                    page_ranges=None,
                    poll_interval=poll_interval,
                    timeout=timeout,
                )
                return index, page_range, markdown
            except Exception as exc:
                raise RuntimeError(f"MinerU parallel page chunk {page_range} failed: {exc}") from exc

        results: list[tuple[int, str, str]] = []
        with ThreadPoolExecutor(max_workers=min(max_workers, len(chunk_jobs))) as executor:
            futures = [
                executor.submit(run_chunk, index, page_range, chunk_path)
                for index, page_range, chunk_path in chunk_jobs
            ]
            for future in as_completed(futures):
                results.append(future.result())

    merged = []
    for _, page_range, markdown in sorted(results, key=lambda item: item[0]):
        merged.append(f"## MinerU 页段 {page_range}\n\n{markdown.strip()}")
    return "\n\n".join(merged).strip()


def _get_pdf_page_count(path: Path) -> int:
    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("Please install pdfplumber before using parallel MinerU parsing.") from exc
    with pdfplumber.open(path) as pdf:
        return len(pdf.pages)


def _write_pdf_page_chunk(source_pdf: Path, pages: List[int], output_pdf: Path) -> None:
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        try:
            from PyPDF2 import PdfReader, PdfWriter
        except ImportError as exc:
            raise RuntimeError(
                "Please install pypdf before using parse_method=mineru_parallel_pages. "
                "Run: pip install pypdf"
            ) from exc

    reader = PdfReader(str(source_pdf))
    writer = PdfWriter()
    for page in pages:
        writer.add_page(reader.pages[page - 1])

    output_pdf.parent.mkdir(parents=True, exist_ok=True)
    with output_pdf.open("wb") as file_obj:
        writer.write(file_obj)


def _chunk_contiguous_pages(pages: List[int], chunk_size: int) -> List[List[int]]:
    chunks: List[List[int]] = []
    current: List[int] = []
    for page in pages:
        if not current:
            current = [page]
            continue
        if page == current[-1] + 1 and len(current) < chunk_size:
            current.append(page)
            continue
        chunks.append(current)
        current = [page]
    if current:
        chunks.append(current)
    return chunks


def _format_page_range(pages: List[int]) -> str:
    if not pages:
        return ""
    if len(pages) == 1:
        return str(pages[0])
    return f"{pages[0]}-{pages[-1]}"


def _pdf_has_text_layer(path: Path, sample_pages: int = 5) -> bool:
    if not path.exists():
        return False
    try:
        import pdfplumber
    except ImportError:
        return False
    try:
        with pdfplumber.open(path) as pdf:
            pages = pdf.pages[: max(1, sample_pages)]
            if not pages:
                return False
            chars = 0
            for page in pages:
                chars += len((page.extract_text() or "").strip())
            return chars >= 20
    except Exception:
        return False


def _pdf_has_tables(path: Path, sample_pages: int = 5) -> bool:
    if not path.exists():
        return False
    try:
        import pdfplumber
    except ImportError:
        return False
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[: max(1, sample_pages)]:
                try:
                    if page.extract_tables():
                        return True
                except Exception:
                    continue
    except Exception:
        return False
    return False


def _pdf_has_images(path: Path, sample_pages: int = 5) -> bool:
    if not path.exists():
        return False
    try:
        import pdfplumber
    except ImportError:
        return False
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[: max(1, sample_pages)]:
                if getattr(page, "images", None):
                    return True
    except Exception:
        return False
    return False


def _module_available(module_name: str) -> bool:
    return importlib.util.find_spec(module_name) is not None


def parse_bid_document_as_list(
    document: str,
    output_dir: Optional[str] = None,
    parse_method: str = "mineru_vlm",
    model_version: Optional[str] = None,
    language: str = "ch",
    is_ocr: Optional[bool] = None,
    enable_formula: bool = True,
    enable_table: bool = True,
    page_ranges: Optional[str] = None,
    poll_interval: int = 5,
    timeout: int = 600,
) -> List[dict]:
    """Parse a bid document and return plain dicts for JSON/API responses."""
    sections = parse_bid_document(
        document=document,
        output_dir=output_dir,
        parse_method=parse_method,
        model_version=model_version,
        language=language,
        is_ocr=is_ocr,
        enable_formula=enable_formula,
        enable_table=enable_table,
        page_ranges=page_ranges,
        poll_interval=poll_interval,
        timeout=timeout,
    )
    return [section.model_dump() for section in sections]


def split_bid_markdown_sections(markdown: str) -> List[BidDocumentSection]:
    """Split Markdown into chapter/section chunks."""
    lines = markdown.splitlines()
    heading_positions = []

    for line_number, line in enumerate(lines):
        heading = _detect_heading(line)
        if heading:
            level, title = heading
            heading_positions.append((line_number, level, title))

    if not heading_positions:
        content = markdown.strip()
        return [
            BidDocumentSection(
                index=1,
                title="全文",
                level=1,
                content=content,
                markdown=content,
            )
        ]

    sections: List[BidDocumentSection] = []
    title_stack: list[tuple[int, str]] = []
    preface = "\n".join(lines[: heading_positions[0][0]]).strip()
    if preface:
        sections.append(
            BidDocumentSection(
                index=1,
                title="前言",
                level=1,
                content=preface,
                markdown=preface,
                metadata={
                    "title_path": ["前言"],
                    "tables": _extract_markdown_tables_json(preface),
                    "images": [],
                    "page_start": _guess_page_number(preface),
                    "page_end": _guess_page_number(preface),
                },
            )
        )

    for idx, (start_line, level, title) in enumerate(heading_positions):
        end_line = (
            heading_positions[idx + 1][0]
            if idx + 1 < len(heading_positions)
            else len(lines)
        )
        chunk_lines = lines[start_line:end_line]
        markdown_chunk = "\n".join(chunk_lines).strip()
        content = "\n".join(chunk_lines[1:]).strip()
        while title_stack and title_stack[-1][0] >= level:
            title_stack.pop()
        title_stack.append((level, title.strip()))
        title_path = [item[1] for item in title_stack]
        sections.append(
            BidDocumentSection(
                index=len(sections) + 1,
                title=title.strip(),
                level=level,
                content=content,
                markdown=markdown_chunk,
                metadata={
                    "title_path": title_path,
                    "tables": _extract_markdown_tables_json(markdown_chunk),
                    "images": [],
                    "page_start": _guess_page_number(markdown_chunk),
                    "page_end": _guess_page_number(markdown_chunk),
                    "heading_detection": "markdown_or_number_pattern",
                },
            )
        )

    return sections


def enrich_sections_with_metadata(
    sections: List[BidDocumentSection],
    *,
    document: str,
) -> List[BidDocumentSection]:
    path = Path(document)
    image_items = _extract_docx_image_metadata(path) if path.suffix.lower() == ".docx" else []
    if not image_items:
        return sections

    for image in image_items:
        target_index = _match_image_to_section(image, sections)
        if target_index is None:
            continue
        metadata = dict(sections[target_index].metadata or {})
        images = list(metadata.get("images") or [])
        images.append(image)
        metadata["images"] = images
        sections[target_index].metadata = metadata
    return sections


def _match_image_to_section(image: dict[str, Any], sections: List[BidDocumentSection]) -> int | None:
    anchor = str(image.get("nearby_text") or image.get("paragraph_text") or "").strip()
    if anchor:
        anchor_key = re.sub(r"\s+", "", anchor)[:40]
        for index, section in enumerate(sections):
            section_text = re.sub(r"\s+", "", section.markdown or section.content or "")
            if anchor_key and anchor_key in section_text:
                return index
    return len(sections) - 1 if sections else None


def _extract_markdown_tables_json(markdown: str) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    table_lines: list[str] = []
    for line in str(markdown or "").splitlines() + [""]:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines.append(stripped)
            continue
        if table_lines:
            parsed = _parse_markdown_table(table_lines)
            if parsed:
                tables.append(parsed)
            table_lines = []
    return tables


def _parse_markdown_table(lines: list[str]) -> dict[str, Any] | None:
    data_lines = [
        line
        for line in lines
        if not re.match(r"^\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?$", line)
    ]
    if not data_lines:
        return None
    rows = _expand_merged_like_table_rows(
        [[cell.strip() for cell in line.split("|")[1:-1]] for line in data_lines]
    )
    header = rows[0] if rows else []
    body = rows[1:] if len(rows) > 1 else []
    return {
        "headers": header,
        "rows": body,
        "row_cell_header_maps": [
            _build_cell_header_map(row, header)
            for row in body
        ],
        "row_count": len(body),
        "markdown": "\n".join(lines),
    }


def _expand_merged_like_table_rows(rows: list[list[str]]) -> list[list[str]]:
    expanded: list[list[str]] = []
    previous: list[str] = []
    for row_index, row in enumerate(rows):
        normalized = _expand_merged_like_table_row(
            row,
            previous if row_index > 0 else [],
            fill_from_left=row_index == 0,
        )
        expanded.append(normalized)
        previous = normalized
    return expanded


def _expand_merged_like_table_row(
    row: list[str],
    previous_row: list[str] | None = None,
    *,
    fill_from_left: bool = False,
) -> list[str]:
    previous = previous_row or []
    width = max(len(row), len(previous))
    output: list[str] = []
    last_value = ""
    for index in range(width):
        value = str(row[index]).strip() if index < len(row) else ""
        if value:
            last_value = value
        elif fill_from_left and last_value:
            value = last_value
        elif index < len(previous):
            value = str(previous[index]).strip()
        output.append(value)
    return output


def _build_cell_header_map(row: list[str], header: list[str]) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for index, cell in enumerate(row):
        key = str(header[index]).strip() if index < len(header) and str(header[index]).strip() else f"col_{index + 1}"
        mapping[key] = str(cell).strip()
    return mapping


def _normalize_openpyxl_merged_cells(ws: Any) -> None:
    merged_ranges = list(getattr(getattr(ws, "merged_cells", None), "ranges", []) or [])
    for merged_range in merged_ranges:
        min_row, min_col = merged_range.min_row, merged_range.min_col
        max_row, max_col = merged_range.max_row, merged_range.max_col
        value = ws.cell(row=min_row, column=min_col).value
        ws.unmerge_cells(str(merged_range))
        for row in range(min_row, max_row + 1):
            for col in range(min_col, max_col + 1):
                ws.cell(row=row, column=col).value = value


def _guess_page_number(markdown: str) -> int | None:
    match = re.search(r"第\s*(\d+)\s*页", str(markdown or ""))
    return int(match.group(1)) if match else None


def _resolve_mineru_method(
    parse_method: str,
    model_version: Optional[str],
    is_ocr: Optional[bool],
) -> tuple[str, bool]:
    if parse_method == "mineru_pipeline":
        return model_version or "pipeline", bool(is_ocr) if is_ocr is not None else False
    if parse_method == "mineru_html":
        return "MinerU-HTML", False
    return model_version or "vlm", bool(is_ocr) if is_ocr is not None else False


def _parse_pdf_with_pdfplumber(document: str, page_ranges: Optional[str]) -> str:
    path = Path(document).expanduser().resolve()
    if path.suffix.lower() != ".pdf":
        raise ValueError("pdfplumber only supports local PDF files.")
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("Please install pdfplumber before using parse_method=pdfplumber.") from exc

    selected_pages = _parse_page_ranges(page_ranges)
    page_texts = []
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            if selected_pages and page_index not in selected_pages:
                continue
            text = page.extract_text() or ""
            page_texts.append(f"## 第 {page_index} 页\n{text.strip()}")
    return "\n\n".join(page_texts).strip()


def _parse_docx_with_docx2python(document: str, include_image_ocr: bool = False) -> str:
    path = Path(document).expanduser().resolve()
    if path.suffix.lower() == ".doc":
        with TemporaryDirectory() as temp_dir:
            converted_path = Path(temp_dir) / f"{path.stem}.docx"
            _convert_doc_to_docx_with_word(path, converted_path)
            return _parse_docx_with_docx2python(str(converted_path), include_image_ocr=include_image_ocr)
    if path.suffix.lower() != ".docx":
        raise ValueError("docx2python only supports local Word files.")
    if not path.exists():
        raise FileNotFoundError(f"Word file not found: {path}")

    docx2python_text = _extract_docx_with_docx2python(path)
    if _is_enough_word_content(docx2python_text):
        return docx2python_text

    python_docx_text = _extract_docx_with_python_docx(path)
    merged_text = _merge_text_blocks([docx2python_text, python_docx_text])
    if not merged_text:
        raise RuntimeError("Word parse result is empty. Try MinerU Pipeline for this file.")
    if include_image_ocr:
        image_ocr_text = _extract_docx_images_ocr_fast(path)
        if image_ocr_text:
            return f"{merged_text}\n\n## 图片OCR内容\n\n{image_ocr_text}".strip()
    return merged_text


def _is_enough_word_content(text: str) -> bool:
    clean_text = text.strip()
    if len(clean_text) >= 1500:
        return True
    non_empty_lines = [line for line in clean_text.splitlines() if line.strip()]
    return len(non_empty_lines) >= 40


def _convert_doc_to_docx_with_word(source_path: Path, target_path: Path) -> None:
    if not source_path.exists():
        raise FileNotFoundError(f"Word file not found: {source_path}")

    command = [
        "powershell",
        "-NoProfile",
        "-ExecutionPolicy",
        "Bypass",
        "-Command",
        (
            "$ErrorActionPreference = 'Stop'; "
            "$word = New-Object -ComObject Word.Application; "
            "$word.Visible = $false; "
            "try { "
            f"$doc = $word.Documents.Open('{_escape_powershell_single_quote(source_path)}'); "
            f"$doc.SaveAs([ref] '{_escape_powershell_single_quote(target_path)}', [ref] 16); "
            "$doc.Close([ref] $false); "
            "} finally { $word.Quit() }"
        ),
    ]
    result = subprocess.run(
        command,
        capture_output=True,
        text=True,
        encoding="utf-8",
        errors="replace",
        timeout=180,
        check=False,
    )
    if result.returncode != 0 or not target_path.exists():
        raise RuntimeError(
            "Failed to convert .doc to .docx with Microsoft Word. "
            "Please save the file as .docx manually or use MinerU. "
            f"stderr={result.stderr[-1000:]}"
        )


def _escape_powershell_single_quote(path: Path) -> str:
    return str(path).replace("'", "''")


def _extract_docx_with_docx2python(path: Path) -> str:
    try:
        from docx2python import docx2python
    except ImportError as exc:
        raise RuntimeError("Please install docx2python before using parse_method=docx2python.") from exc

    with docx2python(path) as docx_content:
        parts = [
            getattr(docx_content, "text", ""),
            _flatten_docx2python_content(getattr(docx_content, "body", "")),
            _flatten_docx2python_content(getattr(docx_content, "header", "")),
            _flatten_docx2python_content(getattr(docx_content, "footer", "")),
            _flatten_docx2python_content(getattr(docx_content, "footnotes", "")),
            _flatten_docx2python_content(getattr(docx_content, "endnotes", "")),
        ]
    return _merge_text_blocks(parts)


def _extract_docx_with_python_docx(path: Path) -> str:
    try:
        from docx import Document
    except ImportError:
        return ""

    document = Document(path)
    lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            lines.append(text)

    for table in document.tables:
        lines.extend(_extract_docx_table_lines(table))

    for section in document.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in section.header.tables:
            lines.extend(_extract_docx_table_lines(table))
        for paragraph in section.footer.paragraphs:
            text = paragraph.text.strip()
            if text:
                lines.append(text)
        for table in section.footer.tables:
            lines.extend(_extract_docx_table_lines(table))

    return "\n".join(lines).strip()


def _extract_docx_images_ocr(path: Path) -> str:
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return "本地 OCR Python 依赖未安装，未执行图片文字识别。"

    image_blocks = []
    with TemporaryDirectory() as temp_dir:
        image_paths = _extract_docx_images(path, Path(temp_dir))
        for index, image_path in enumerate(image_paths, start=1):
            try:
                with Image.open(image_path) as image:
                    text = pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
            except Exception as exc:
                text = f"OCR识别失败：{exc}"
            image_blocks.append(
                "\n".join(
                    [
                        f"### 图片 {index}",
                        f"- 文件：{image_path.name}",
                        "- 来源：Word 内嵌图片",
                        "",
                        text or "未识别到文字。",
                    ]
                )
            )
    return "\n\n".join(image_blocks).strip()


def _extract_docx_images(path: Path, output_dir: Path) -> List[Path]:
    image_paths = []
    with zipfile.ZipFile(path) as archive:
        image_names = [
            name
            for name in archive.namelist()
            if name.startswith("word/media/") and not name.endswith("/")
        ]
        for index, name in enumerate(image_names, start=1):
            suffix = Path(name).suffix or ".png"
            target_path = output_dir / f"image_{index:03d}{suffix}"
            target_path.write_bytes(archive.read(name))
            image_paths.append(target_path)
    return image_paths


def _extract_docx_image_metadata(path: Path) -> list[dict[str, Any]]:
    if not path.exists() or path.suffix.lower() != ".docx":
        return []
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        return []

    try:
        document = Document(path)
    except Exception:
        return []

    relationship_to_name = _docx_image_relationship_names(path)
    output: list[dict[str, Any]] = []

    paragraph_context = []
    for paragraph_index, paragraph in enumerate(document.paragraphs, start=1):
        paragraph_context.append((paragraph_index, paragraph.text.strip(), paragraph))

    for paragraph_index, text, paragraph in paragraph_context:
        rel_ids = _paragraph_image_rel_ids(paragraph)
        if not rel_ids:
            continue
        nearby = _nearby_paragraph_text(paragraph_context, paragraph_index)
        for rel_id in rel_ids:
            file_name = relationship_to_name.get(rel_id, rel_id)
            output.append(
                _build_docx_image_metadata(
                    image_index=len(output) + 1,
                    source=file_name,
                    location_type="paragraph",
                    paragraph_index=paragraph_index,
                    table_index=None,
                    nearby_text=nearby,
                    paragraph_text=text,
                    path=path,
                    rel_id=rel_id,
                )
            )

    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            for cell_index, cell in enumerate(row.cells, start=1):
                cell_text = " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                for paragraph in cell.paragraphs:
                    rel_ids = _paragraph_image_rel_ids(paragraph)
                    if not rel_ids:
                        continue
                    for rel_id in rel_ids:
                        file_name = relationship_to_name.get(rel_id, rel_id)
                        output.append(
                            _build_docx_image_metadata(
                                image_index=len(output) + 1,
                                source=file_name,
                                location_type="table_cell",
                                paragraph_index=row_index,
                                table_index=table_index,
                                nearby_text=cell_text,
                                paragraph_text=paragraph.text.strip(),
                                path=path,
                                rel_id=rel_id,
                                extra={
                                    "row_index": row_index,
                                    "cell_index": cell_index,
                                },
                            )
                        )
    return output


def _docx_image_relationship_names(path: Path) -> dict[str, str]:
    names: dict[str, str] = {}
    try:
        with zipfile.ZipFile(path) as archive:
            rels_name = "word/_rels/document.xml.rels"
            if rels_name not in archive.namelist():
                return names
            rels_xml = archive.read(rels_name).decode("utf-8", errors="ignore")
            for match in re.finditer(
                r'Id="([^"]+)".+?Type="[^"]+/image".+?Target="([^"]+)"',
                rels_xml,
            ):
                names[match.group(1)] = Path(match.group(2)).name
    except Exception:
        pass
    return names


def _paragraph_image_rel_ids(paragraph: Any) -> list[str]:
    rel_ids: list[str] = []
    try:
        for blip in paragraph._p.iter():
            embed = blip.attrib.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
            if embed:
                rel_ids.append(embed)
    except Exception:
        return rel_ids
    return rel_ids


def _nearby_paragraph_text(paragraph_context: list[tuple[int, str, Any]], paragraph_index: int) -> str:
    texts = [
        text
        for index, text, _ in paragraph_context
        if abs(index - paragraph_index) <= 2 and text
    ]
    return " / ".join(texts)


def _build_docx_image_metadata(
    *,
    image_index: int,
    source: str,
    location_type: str,
    paragraph_index: int,
    table_index: int | None,
    nearby_text: str,
    paragraph_text: str,
    path: Path,
    rel_id: str,
    extra: dict[str, Any] | None = None,
) -> dict[str, Any]:
    ocr_text = _ocr_docx_relationship_image(path, rel_id)
    visual_guess = _guess_visual_content_type(ocr_text, nearby_text)
    metadata = {
        "index": image_index,
        "source": source,
        "location_type": location_type,
        "paragraph_index": paragraph_index,
        "table_index": table_index,
        "paragraph_text": paragraph_text,
        "nearby_text": nearby_text,
        "ocr_text": ocr_text,
        "visual_guess": visual_guess,
        "vlm_status": "pending_external_vlm",
    }
    if extra:
        metadata.update(extra)
    return metadata


def _ocr_docx_relationship_image(path: Path, rel_id: str) -> str:
    rel_names = _docx_image_relationship_names(path)
    file_name = rel_names.get(rel_id)
    if not file_name:
        return ""
    try:
        with TemporaryDirectory() as temp_dir, zipfile.ZipFile(path) as archive:
            candidates = [
                name
                for name in archive.namelist()
                if name.startswith("word/media/") and Path(name).name == file_name
            ]
            if not candidates:
                return ""
            target = Path(temp_dir) / file_name
            target.write_bytes(archive.read(candidates[0]))
            return _ocr_image_with_local_engines(target)
    except Exception:
        return ""


def _guess_visual_content_type(ocr_text: str, nearby_text: str) -> str:
    text = f"{ocr_text}\n{nearby_text}"
    if re.search(r"公章|盖章|印章|有限公司|委员会|中心|局", text):
        return "疑似公章/印章"
    if re.search(r"签字|签名|法定代表人|授权代表|日期|年\s*月\s*日", text):
        return "疑似签字/签署区域"
    if re.search(r"流程|架构|示意|图|网络|系统", text):
        return "疑似图表/架构图"
    return "普通图片"


def _extract_docx_images_ocr_fast(path: Path) -> str:
    image_blocks = []
    with TemporaryDirectory() as temp_dir:
        image_paths = _extract_docx_images(path, Path(temp_dir))
        for index, image_path in enumerate(image_paths, start=1):
            text = _ocr_image_with_local_engines(image_path)
            image_blocks.append(
                "\n".join(
                    [
                        f"### 图片 {index}",
                        f"- 文件：{image_path.name}",
                        "- 来源：Word 内嵌图片",
                        "",
                        text or "未识别到文字。",
                    ]
                )
            )
    return "\n\n".join(image_blocks).strip()


def _ocr_image_with_local_engines(image_path: Path) -> str:
    rapidocr_text = _ocr_image_with_rapidocr(image_path)
    if rapidocr_text:
        return rapidocr_text

    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return "本地 OCR Python 依赖未安装，未执行图片文字识别。"

    try:
        with Image.open(image_path) as image:
            return pytesseract.image_to_string(image, lang="chi_sim+eng").strip()
    except Exception as exc:
        return f"OCR识别失败：{exc}。建议安装 RapidOCR 或 Tesseract OCR 引擎。"


def _ocr_image_with_rapidocr(image_path: Path) -> str:
    try:
        from rapidocr_onnxruntime import RapidOCR
    except ImportError:
        return ""

    try:
        engine = RapidOCR()
        result, _ = engine(str(image_path))
        if not result:
            return ""
        return "\n".join(item[1] for item in result if len(item) > 1 and item[1]).strip()
    except Exception:
        return ""


def _extract_docx_table_lines(table: Any) -> List[str]:
    lines = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            cell_text = " ".join(
                paragraph.text.strip()
                for paragraph in cell.paragraphs
                if paragraph.text.strip()
            )
            if cell_text:
                cells.append(cell_text)
        if cells:
            lines.append(" | ".join(cells))
    return lines


def _flatten_docx2python_content(value: Any) -> str:
    lines = []

    def walk(item: Any) -> None:
        if isinstance(item, str):
            text = item.strip()
            if text:
                lines.append(text)
            return
        if isinstance(item, dict):
            for nested in item.values():
                walk(nested)
            return
        if isinstance(item, (list, tuple)):
            for nested in item:
                walk(nested)

    walk(value)
    return "\n".join(lines).strip()


def _merge_text_blocks(blocks: List[str]) -> str:
    seen = set()
    lines = []
    for block in blocks:
        for line in block.splitlines():
            clean_line = re.sub(r"\s+", " ", line).strip()
            if not clean_line:
                continue
            if clean_line in seen:
                continue
            seen.add(clean_line)
            lines.append(clean_line)
    return "\n".join(lines).strip()


def _parse_pdf_with_pymupdf4llm(document: str, page_ranges: Optional[str]) -> str:
    path = Path(document).expanduser().resolve()
    if path.suffix.lower() != ".pdf":
        raise ValueError("PyMuPDF4LLM only supports local PDF files.")
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    with _pdf_source_for_page_ranges(path, page_ranges) as source_pdf:
        try:
            import pymupdf4llm

            markdown = pymupdf4llm.to_markdown(str(source_pdf)) or ""
            if markdown.strip():
                return markdown.strip()
        except ImportError:
            pass
        except Exception:
            pass

        return _parse_pdf_with_pymupdf_basic(str(source_pdf), page_ranges=None)


def _parse_pdf_with_pymupdf_basic(document: str, page_ranges: Optional[str]) -> str:
    path = Path(document).expanduser().resolve()
    try:
        import fitz
    except ImportError as exc:
        raise RuntimeError(
            "Please install pymupdf4llm or pymupdf before using parse_method=pymupdf4llm."
        ) from exc

    selected_pages = _parse_page_ranges(page_ranges)
    page_blocks = []
    with fitz.open(str(path)) as pdf:
        for page_index in range(pdf.page_count):
            page_number = page_index + 1
            if selected_pages and page_number not in selected_pages:
                continue
            page = pdf.load_page(page_index)
            text = page.get_text("text") or ""
            blocks = [f"<!-- page:{page_number} -->", f"## 第 {page_number} 页"]
            if text.strip():
                blocks.append(text.strip())
            image_count = len(page.get_images(full=True) or [])
            if image_count:
                blocks.append(f"### 图片说明\n本页检测到 {image_count} 个图片对象，建议结合图片解析模块核对。")
            page_blocks.append("\n\n".join(blocks))
    return "\n\n".join(page_blocks).strip()


def _parse_pdf_with_docling(document: str, page_ranges: Optional[str]) -> str:
    path = Path(document).expanduser().resolve()
    if path.suffix.lower() != ".pdf":
        raise ValueError("Docling only supports local PDF files.")
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    try:
        from docling.document_converter import DocumentConverter
    except ImportError as exc:
        raise RuntimeError(
            "Please install docling before using parse_method=docling. Run: pip install docling"
        ) from exc

    with _pdf_source_for_page_ranges(path, page_ranges) as source_pdf:
        try:
            result = DocumentConverter().convert(str(source_pdf))
            markdown = result.document.export_to_markdown() or ""
        except Exception as exc:
            raise RuntimeError(f"Docling parse failed: {exc}") from exc
    return markdown.strip()


@contextmanager
def _pdf_source_for_page_ranges(path: Path, page_ranges: Optional[str]):
    selected_pages = sorted(_parse_page_ranges(page_ranges))
    if not selected_pages:
        yield path
        return

    page_count = _get_pdf_page_count(path)
    out_of_range_pages = [page for page in selected_pages if page < 1 or page > page_count]
    if out_of_range_pages:
        raise RuntimeError(
            f"Page range contains pages outside the PDF page count ({page_count}): "
            f"{_format_page_range(out_of_range_pages)}"
        )

    with TemporaryDirectory() as temp_dir:
        selected_pdf = Path(temp_dir) / f"{path.stem}_selected_pages.pdf"
        _write_pdf_page_chunk(source_pdf=path, pages=selected_pages, output_pdf=selected_pdf)
        yield selected_pdf


def _parse_pdf_with_pdfplumber(document: str, page_ranges: Optional[str]) -> str:
    path = Path(document).expanduser().resolve()
    if path.suffix.lower() != ".pdf":
        raise ValueError("pdfplumber only supports local PDF files.")
    if not path.exists():
        raise FileNotFoundError(f"PDF file not found: {path}")

    try:
        import pdfplumber
    except ImportError as exc:
        raise RuntimeError("Please install pdfplumber before using parse_method=pdfplumber.") from exc

    selected_pages = _parse_page_ranges(page_ranges)
    page_texts = []
    with pdfplumber.open(path) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            if selected_pages and page_index not in selected_pages:
                continue
            text = page.extract_text() or ""
            blocks = [f"<!-- page:{page_index} -->", f"## 第 {page_index} 页"]
            if text.strip():
                blocks.append(text.strip())
            table_blocks = _extract_pdfplumber_tables_as_markdown(page)
            if table_blocks:
                blocks.append("### 表格")
                blocks.extend(table_blocks)
            if getattr(page, "images", None):
                blocks.append(f"### 图片说明\n本页检测到 {len(page.images)} 个图片/图形对象，建议结合图片解析模块核对。")
            page_texts.append("\n\n".join(blocks))
    return "\n\n".join(page_texts).strip()


def _extract_pdfplumber_tables_as_markdown(page: Any) -> List[str]:
    try:
        tables = page.extract_tables() or []
    except Exception:
        return []

    markdown_tables = []
    for table in tables:
        cleaned_rows = [
            [str(cell or "").strip().replace("\n", " ") for cell in row]
            for row in table
            if row and any(str(cell or "").strip() for cell in row)
        ]
        if not cleaned_rows:
            continue
        width = max(len(row) for row in cleaned_rows)
        normalized = [row + [""] * (width - len(row)) for row in cleaned_rows]
        header = normalized[0]
        body = normalized[1:] or [[""] * width]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join("---" for _ in header) + " |",
        ]
        lines.extend("| " + " | ".join(row) + " |" for row in body)
        markdown_tables.append("\n".join(lines))
    return markdown_tables


def _parse_page_ranges(page_ranges: Optional[str]) -> set[int]:
    if not page_ranges:
        return set()

    pages: set[int] = set()
    for item in page_ranges.split(","):
        item = item.strip()
        if not item:
            continue
        if "-" in item:
            start, end = item.split("-", 1)
            pages.update(range(int(start), int(end) + 1))
        else:
            pages.add(int(item))
    return pages


def _detect_heading(line: str) -> Optional[tuple[int, str]]:
    stripped = line.strip()
    if not stripped:
        return None

    markdown_match = MARKDOWN_HEADING_RE.match(stripped)
    if markdown_match:
        return len(markdown_match.group(1)), markdown_match.group(2).strip()

    for pattern, level in (
        (CHINESE_CHAPTER_RE, 1),
        (CHINESE_ORDER_RE, 2),
        (NUMBERED_HEADING_RE, 2),
        (PAREN_HEADING_RE, 3),
    ):
        match = pattern.match(stripped)
        if match and _looks_like_heading(stripped):
            return level, match.group(1).strip()

    return None


def _looks_like_heading(text: str) -> bool:
    if len(text) > 80:
        return False
    if text.endswith(("。", "，", ",", "；", ";")):
        return False
    return True
